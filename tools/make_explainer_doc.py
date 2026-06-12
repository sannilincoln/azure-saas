"""
Generates a newbie-friendly Word (.docx) explainer for the BadgeMeUp / Azure SaaS app,
with hand-drawn diagrams (Pillow) embedded. No pandoc/LibreOffice needed.

Run:  python tools/make_explainer_doc.py
Out:  How-BadgeMeUp-Works-Explained.docx  (repo root)
"""
import os
from PIL import Image, ImageDraw, ImageFont

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
IMG_DIR = os.path.join(ROOT, "tools", "_diagrams")
os.makedirs(IMG_DIR, exist_ok=True)

FONT = "/c/Windows/Fonts/arial.ttf".replace("/c/", "C:/")
FONT_B = "/c/Windows/Fonts/arialbd.ttf".replace("/c/", "C:/")

# ---- colour palette (soft, friendly) ----
BLUE = (210, 229, 247)
BLUE_E = (90, 140, 200)
GREEN = (212, 237, 218)
GREEN_E = (90, 170, 110)
ORANGE = (255, 234, 204)
ORANGE_E = (220, 160, 70)
PURPLE = (228, 219, 245)
PURPLE_E = (150, 110, 200)
GREY = (235, 235, 235)
GREY_E = (170, 170, 170)
TXT = (30, 30, 30)
ARROW = (90, 90, 90)


def font(sz, bold=False):
    return ImageFont.truetype(FONT_B if bold else FONT, sz)


def wrap(draw, text, fnt, max_w):
    words, lines, cur = text.split(), [], ""
    for w in words:
        t = (cur + " " + w).strip()
        if draw.textlength(t, font=fnt) <= max_w:
            cur = t
        else:
            if cur:
                lines.append(cur)
            cur = w
    if cur:
        lines.append(cur)
    return lines


def box(draw, x, y, w, h, title, sub=None, fill=BLUE, edge=BLUE_E):
    draw.rounded_rectangle([x, y, x + w, y + h], radius=12, fill=fill, outline=edge, width=3)
    tf = font(21, bold=True)
    lines = wrap(draw, title, tf, w - 24)
    sf = font(17)
    sub_lines = wrap(draw, sub, sf, w - 24) if sub else []
    total_h = len(lines) * 26 + (len(sub_lines) * 21 + 6 if sub_lines else 0)
    cy = y + (h - total_h) / 2
    for ln in lines:
        tw = draw.textlength(ln, font=tf)
        draw.text((x + (w - tw) / 2, cy), ln, font=tf, fill=TXT)
        cy += 26
    if sub_lines:
        cy += 6
        for ln in sub_lines:
            tw = draw.textlength(ln, font=sf)
            draw.text((x + (w - tw) / 2, cy), ln, font=sf, fill=(80, 80, 80))
            cy += 21


def arrow(draw, x1, y1, x2, y2, label=None, color=ARROW):
    draw.line([x1, y1, x2, y2], fill=color, width=3)
    import math
    ang = math.atan2(y2 - y1, x2 - x1)
    L = 13
    for da in (math.radians(150), math.radians(-150)):
        draw.line([x2, y2, x2 + L * math.cos(ang + da), y2 + L * math.sin(ang + da)], fill=color, width=3)
    if label:
        lf = font(15)
        mx, my = (x1 + x2) / 2, (y1 + y2) / 2
        tw = draw.textlength(label, font=lf)
        draw.rectangle([mx - tw / 2 - 4, my - 12, mx + tw / 2 + 4, my + 10], fill=(255, 255, 255))
        draw.text((mx - tw / 2, my - 10), label, font=lf, fill=(70, 70, 70))


def canvas(w, h):
    img = Image.new("RGB", (w, h), (255, 255, 255))
    return img, ImageDraw.Draw(img)


# ============ Diagram 1: The big picture (architecture) ============
def diagram_architecture():
    img, d = canvas(1100, 760)
    box(d, 430, 30, 240, 70, "You", "(your web browser)", GREY, GREY_E)

    box(d, 120, 170, 300, 90, "Sign-up & Admin Website", "where new orgs sign up", BLUE, BLUE_E)
    box(d, 680, 170, 300, 90, "BadgeMeUp Website", "the main app you log into", BLUE, BLUE_E)

    box(d, 442, 172, 216, 86, "Microsoft Entra External ID",
        "the ID checkpoint (login)", PURPLE, PURPLE_E)

    box(d, 120, 360, 300, 90, "Admin API", "the rules & tenant brain", GREEN, GREEN_E)
    box(d, 680, 360, 300, 90, "Permissions API", "who-can-do-what service", GREEN, GREEN_E)

    box(d, 120, 540, 300, 80, "Tenant Database", "list of organisations", ORANGE, ORANGE_E)
    box(d, 680, 540, 300, 80, "Permissions Database", "each user's access rights", ORANGE, ORANGE_E)

    box(d, 410, 545, 280, 75, "Key Vault + App Config", "secrets & settings", GREY, GREY_E)

    # arrows
    arrow(d, 480, 100, 300, 170, "sign up")
    arrow(d, 620, 100, 800, 170, "log in")
    arrow(d, 550, 100, 550, 175, "")
    arrow(d, 270, 260, 270, 360, "asks")
    arrow(d, 830, 260, 830, 360, "")
    arrow(d, 420, 405, 680, 405, "checks")
    arrow(d, 270, 450, 270, 540, "reads")
    arrow(d, 830, 450, 830, 540, "reads")
    arrow(d, 550, 450, 550, 545, "")
    img.save(os.path.join(IMG_DIR, "architecture.png"))


# ============ Diagram 2: Signing in ============
def diagram_signin():
    img, d = canvas(1100, 430)
    y = 150
    box(d, 30, y, 200, 90, "1. You", "click 'Log in'", GREY, GREY_E)
    box(d, 300, y, 230, 90, "2. Entra External ID", "asks for email + password", PURPLE, PURPLE_E)
    box(d, 600, y, 220, 90, "3. Digital pass", "(a signed token) is issued", BLUE, BLUE_E)
    box(d, 890, y, 180, 90, "4. App opens", "you're in!", GREEN, GREEN_E)
    arrow(d, 230, y + 45, 300, y + 45)
    arrow(d, 530, y + 45, 600, y + 45)
    arrow(d, 820, y + 45, 890, y + 45)
    note = font(17)
    d.text((300, 60), "The app never sees your password — only Microsoft does.", font=note, fill=(120, 80, 80))
    d.text((300, 300), "The 'pass' proves who you are for the rest of your visit.", font=note, fill=(80, 80, 120))
    img.save(os.path.join(IMG_DIR, "signin.png"))


# ============ Diagram 3: Permissions (what you're allowed to do) ============
def diagram_permissions():
    img, d = canvas(1180, 360)
    y = 130
    box(d, 20, y, 250, 100, "1. You open a page", "e.g. /sannicafe", GREY, GREY_E)
    box(d, 320, y, 270, 100, "2. Admin API", "'who is this, and are they allowed?'", GREEN, GREEN_E)
    box(d, 640, y, 290, 100, "3. Permissions service", "checks the database: you are Admin of Sanni cafe", ORANGE, ORANGE_E)
    box(d, 980, y, 180, 100, "4. Page loads", "with real data", BLUE, BLUE_E)
    arrow(d, 270, y + 50, 320, y + 50)
    arrow(d, 590, y + 50, 640, y + 50, "asks")
    arrow(d, 930, y + 50, 980, y + 50, "allowed")
    d.text((20, 40), "Logging in proves WHO you are. This separate check decides WHAT you may do.",
           font=font(19, bold=True), fill=TXT)
    note = font(17)
    d.text((20, 290), "Your rights are checked live on the server every time — not baked into your login pass, "
                      "so access changes take effect almost instantly.",
           font=note, fill=(80, 80, 80))
    img.save(os.path.join(IMG_DIR, "permissions.png"))


# ============ Diagram 4: New organisation onboarding ============
def diagram_onboarding():
    img, d = canvas(1100, 300)
    steps = [
        ("1. Sign up", "create your login", PURPLE, PURPLE_E),
        ("2. Name your org", "e.g. 'Sanni cafe'", BLUE, BLUE_E),
        ("3. Pick a web address", "/sannicafe", BLUE, BLUE_E),
        ("4. System sets it up", "creates your space", GREEN, GREEN_E),
        ("5. Ready!", "your org is live", GREEN, GREEN_E),
    ]
    x = 20
    for i, (t, s, f, e) in enumerate(steps):
        box(d, x, 110, 185, 90, t, s, f, e)
        if i < len(steps) - 1:
            arrow(d, x + 185, 155, x + 205, 155)
        x += 205
    d.text((20, 40), "Onboarding: turning a brand-new sign-up into a working organisation",
           font=font(19, bold=True), fill=TXT)
    img.save(os.path.join(IMG_DIR, "onboarding.png"))


for fn in (diagram_architecture, diagram_signin, diagram_permissions, diagram_onboarding):
    fn()
print("diagrams done")

# =========================================================================
# Build the Word document
# =========================================================================
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# base style
st = doc.styles["Normal"]
st.font.name = "Calibri"
st.font.size = Pt(11.5)


def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def para(text, italic=False, size=11.5):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.font.size = Pt(size)
    return p


def bullet(text):
    doc.add_paragraph(text, style="List Bullet")


def figure(path, caption, width=6.3):
    doc.add_picture(path, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x70, 0x70, 0x70)


# ---- Title ----
t = doc.add_heading("How the BadgeMeUp App Works", 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("A plain-English guide for newcomers")
r.italic = True
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
d2 = doc.add_paragraph()
d2.alignment = WD_ALIGN_PARAGRAPH.CENTER
d2.add_run("Built on Microsoft Azure + Entra External ID").font.size = Pt(11)
doc.add_paragraph()

# ---- 1. What is this? ----
heading("1. What is this app, in one minute?", 1)
para("BadgeMeUp is a 'Software as a Service' (SaaS) application. That's just a fancy way "
     "of saying: it's software you use through a web browser, and many different "
     "organisations share the same app at the same time — without ever seeing each "
     "other's data.")
para("Think of an apartment building. It's one building (one app), but each family has "
     "their own locked apartment (their own private space and data). In our app, each "
     "of those 'apartments' is called a tenant — usually a company or organisation, "
     "like a coffee shop called 'Sanni cafe'.")
para("This project is a starter kit (originally Microsoft's 'Azure SaaS Dev Kit') that "
     "gives you all the hard parts already built: signing up, logging in, keeping each "
     "organisation separate, and controlling who is allowed to do what.")

# ---- 2. Key words ----
heading("2. A few key words (no jargon, promise)", 1)
bullet("Tenant — one organisation using the app (one 'apartment'). Example: Sanni cafe.")
bullet("User — a person who logs in. A user can belong to one or more tenants.")
bullet("Sign-up / Onboarding — the process of creating a brand-new organisation in the app.")
bullet("Identity provider — the trusted 'bouncer' that checks your email and password. "
       "We use Microsoft Entra External ID for this. The app itself never sees your password.")
bullet("Token (a 'digital pass') — after you log in, you're given a signed digital pass "
       "that proves who you are for the rest of your visit.")
bullet("Permissions — the list of things you're allowed to do (for example, 'Admin of "
       "Sanni cafe'). The app checks this before showing you sensitive pages.")

# ---- 3. The big picture ----
heading("3. The big picture: the moving parts", 1)
para("The app isn't one single program — it's a small team of programs that each do one "
     "job and talk to each other. Here's the whole team:")
figure(os.path.join(IMG_DIR, "architecture.png"),
       "Figure 1 — How the pieces fit together.")
para("In words:")
bullet("Sign-up & Admin Website — where a new organisation registers and where admins "
       "manage their organisation.")
bullet("BadgeMeUp Website — the main app that everyday users log into.")
bullet("Microsoft Entra External ID — the login checkpoint. It checks your identity and "
       "hands out the digital pass.")
bullet("Admin API — the 'brain'. The websites ask it questions like 'what is this tenant?' "
       "and 'is this person allowed?'.")
bullet("Permissions API — a specialist service that knows each user's access rights.")
bullet("Databases — one stores the list of organisations (tenants), another stores each "
       "user's permissions.")
bullet("Key Vault & App Configuration — secure lockboxes for passwords/secrets and for "
       "the app's settings, so secrets never sit in the code.")

# ---- 4. Logging in ----
heading("4. What happens when you log in", 1)
para("You never type your password into the app itself. Instead, the app sends you to "
     "Microsoft's secure login page. This is safer — the app only ever receives a "
     "signed 'pass' that says 'yes, this really is them'.")
figure(os.path.join(IMG_DIR, "signin.png"),
       "Figure 2 — Logging in, step by step.")
para("That digital pass travels with you while you click around, so you don't have to "
     "log in again on every page.")

# ---- 5. Permissions ----
heading("5. How the app knows what you're allowed to do", 1)
para("Logging in proves who you are. It does not, by itself, say what you're allowed to "
     "do. Those are two different questions. When you open a sensitive page — say your "
     "organisation's page at /sannicafe — the app quietly checks your permissions first.")
figure(os.path.join(IMG_DIR, "permissions.png"),
       "Figure 3 — Checking your access rights, live, on every request.")
para("Important detail (and the heart of this project's recent work): your permissions "
     "are looked up fresh on the server every single time, rather than being stamped "
     "into your login pass. This is more secure — if your access is removed, it stops "
     "working almost immediately, and your rights never travel around inside the pass "
     "where they could leak or go stale.")

# ---- 6. Onboarding ----
heading("6. Starting a brand-new organisation", 1)
para("When someone brings a new organisation to the app, they go through a short wizard. "
     "By the end, the system has carved out a private space for that organisation.")
figure(os.path.join(IMG_DIR, "onboarding.png"),
       "Figure 4 — From sign-up to a live organisation.")

# ---- 7. Where it all lives ----
heading("7. Where does all this actually run?", 1)
para("Everything lives in Microsoft Azure, the cloud platform. You don't need your own "
     "servers. Each of the websites and APIs runs as an 'App Service' (a managed place "
     "to run a web app). The databases are Azure SQL. Secrets live in Key Vault, "
     "settings in App Configuration, and logins are handled by Entra External ID.")
para("When developers update the app, an automated pipeline (currently GitHub Actions) "
     "builds the new version and publishes it to Azure — no manual copying of files.")

# ---- 8. A note on the recent change ----
heading("8. One recent behind-the-scenes change", 1)
para("This app used to use an older Microsoft login system called 'Azure AD B2C'. "
     "Microsoft stopped offering B2C to new customers, so the app was moved over to "
     "the newer 'Microsoft Entra External ID'. To you as a user, nothing looks "
     "different — you still just log in. Under the hood, the team rewired the login "
     "checkpoint and rebuilt the permissions check (Section 5) to work with the new "
     "system. That migration is now complete and working.")

# ---- Glossary ----
heading("Glossary (quick reference)", 1)
gloss = [
    ("SaaS", "Software you use over the internet, shared by many organisations."),
    ("Tenant", "One organisation's private space in the app."),
    ("Entra External ID", "Microsoft's modern login service used here."),
    ("Token / pass", "A signed proof of who you are, given out after login."),
    ("API", "A behind-the-scenes service that answers questions for the websites."),
    ("Permissions", "The list of actions you're allowed to perform."),
    ("Key Vault", "A secure lockbox in Azure for passwords and secrets."),
    ("App Service", "A managed home in Azure where a website or API runs."),
]
for term, desc in gloss:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(term + " — ")
    r.bold = True
    p.add_run(desc)

out = os.path.join(ROOT, "How-BadgeMeUp-Works-Explained.docx")
doc.save(out)
print("saved:", out)
